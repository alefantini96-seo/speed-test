"""
Test del contratto fra browser e server nella versione online.

L'endpoint restituisce una versione ridotta dei fatti: il browser deve rimandarla
indietro per generare il Word, e il corpo di una richiesta Vercel non puo' superare
i 4,5 MB. Se qualcuno aggiunge un campo al report senza aggiungerlo alla riduzione,
il documento esce mutilo: questi test lo intercettano.
"""
import io
import json
from datetime import date
from pathlib import Path

import pytest
from docx import Document

from speed.core import consenso, diagnose, extract, thirdparty
from speed.io import crux, render, render_docx
from speed.web import fatti_essenziali, serializza, terze_essenziali, valida_url

RADICE = Path(__file__).resolve().parent.parent
FIXTURES = RADICE / "fixtures"
URL = "https://www.bbc.com/"
PROPRI = ["bbci.co.uk"]


@pytest.fixture(scope="module")
def pagina():
    """Compone la risposta dell'endpoint dai fixture, senza toccare la rete."""
    misurazioni = [
        extract.estrai(json.loads((FIXTURES / nome).read_text(encoding="utf-8")),
                       URL, "PHONE", PROPRI)
        for nome in ("psi-bbc-mobile-it.json", "psi-bbc-mobile-it-2.json")
    ]
    accordo = consenso.combina(misurazioni)
    fatti = accordo.fatti
    campo = crux.leggi_record(json.loads((FIXTURES / "crux-bbc.json").read_text(encoding="utf-8")))
    storico = crux.leggi_storico(
        json.loads((FIXTURES / "crux-history-bbc.json").read_text(encoding="utf-8")))
    riepilogo = thirdparty.riepiloga(fatti.richieste, URL, PROPRI)
    problemi = diagnose.diagnostica(fatti, campo["metriche"], riepilogo, accordo)

    risposta = {
        "template": URL, "url": URL,
        "fatti": fatti_essenziali(fatti),
        "campo": {"livello": "url", "metriche": campo["metriche"],
                  "storico": {"url": URL} | storico},
        "terze_parti": terze_essenziali(riepilogo),
        "problemi": problemi,
        "misurazioni": accordo.ripetizioni, "concordi": accordo.concordi,
        "consenso": accordo.descrizione,
    }
    # Come arriva al browser e come torna indietro: passata da JSON.
    return json.loads(json.dumps(risposta, default=serializza, ensure_ascii=False))


def _esecuzione(pagina):
    return {"cliente": "bbc.com", "sito": URL, "data": date.today().isoformat(),
            "form_factor": "PHONE", "pagine": [pagina]}


def test_il_payload_resta_piccolo(pagina):
    peso = len(json.dumps(pagina, ensure_ascii=False).encode("utf-8"))
    assert peso < 40_000, "oltre questa soglia 40 pagine sfondano il limite di Vercel"


def test_la_riduzione_non_porta_le_liste_pesanti(pagina):
    for scartato in ("richieste", "opportunita", "risparmi", "campo_psi"):
        assert scartato not in pagina["fatti"], f"{scartato} non serve a valle: va scartato"


def test_la_riduzione_conserva_cio_che_serve_al_report(pagina):
    for necessario in ("lcp_fasi", "lcp_elemento_snippet", "performance_score"):
        assert necessario in pagina["fatti"]
    for necessario in ("byte_totali", "byte_terzi", "richieste_totali"):
        assert necessario in pagina["terze_parti"]


def test_le_risorse_colpevoli_sopravvivono_dentro_i_problemi(pagina):
    """Le richieste di rete si scartano, ma i file da sistemare devono restare."""
    con_risorse = [p for p in pagina["problemi"] if p.get("risorse")]
    assert con_risorse, "senza risorse il report non e' azionabile"
    url, misura, _terza = con_risorse[0]["risorse"][0]
    assert url.startswith("http") and misura


def test_il_word_si_genera_dal_payload_ridotto(pagina, tmp_path):
    percorso = render_docx.docx_report(_esecuzione(pagina), tmp_path / "r.docx")
    doc = Document(str(percorso))
    testo = "\n".join(p.text for p in doc.paragraphs)
    for tabella in doc.tables:
        for riga in tabella.rows:
            testo += "\n" + " | ".join(c.text for c in riga.cells)
    for atteso in ("p75 reale", "Dove si perde il tempo dell'LCP", "Interventi",
                   "testo di Lighthouse", "Peso della pagina"):
        assert atteso in testo, f"manca nel Word generato dal payload ridotto: {atteso}"


def test_anche_l_html_si_genera_dal_payload_ridotto(pagina):
    h = render.html_report(_esecuzione(pagina))
    assert "<svg" in h and "testo di Lighthouse" in h


# --- l'app WSGI: instradamento e protezione --------------------------------- #

def _chiama(percorso, metodo="GET", corpo=None):
    from app import app
    dati = json.dumps(corpo or {}).encode("utf-8")
    environ = {
        "PATH_INFO": percorso, "REQUEST_METHOD": metodo,
        "CONTENT_LENGTH": str(len(dati)), "wsgi.input": io.BytesIO(dati),
    }
    catturato = {}

    def avvia(stato, intestazioni):
        catturato["stato"] = stato
        catturato["intestazioni"] = dict(intestazioni)

    corpo_risposta = b"".join(app(environ, avvia))
    return catturato["stato"], catturato["intestazioni"], corpo_risposta


def test_la_pagina_viene_servita():
    stato, intestazioni, corpo = _chiama("/")
    assert stato.startswith("200")
    assert "text/html" in intestazioni["Content-Type"]
    assert b"Analisi velocit" in corpo


def test_percorso_sconosciuto_da_404():
    assert _chiama("/inesistente")[0].startswith("404")


def test_gli_endpoint_rifiutano_il_get():
    assert _chiama("/api/analizza")[0].startswith("405")


def test_url_non_valido_riceve_il_rimedio(monkeypatch):
    monkeypatch.setenv("GOOGLE_API_KEY", "finta")
    stato, _, corpo = _chiama("/api/analizza", "POST", {"url": "non-un-url"})
    assert stato.startswith("400")
    assert "https://" in json.loads(corpo)["rimedio"]


def test_l_app_non_chiede_autenticazione(monkeypatch):
    """Non c'e' password: chiunque abbia il link analizza, e l'unico freno e' il
    limite di richieste per origine. E' una scelta esplicita, e va scoperta da un
    test e non in produzione. Qui l'unico rifiuto e' la chiave Google mancante."""
    monkeypatch.delenv("GOOGLE_API_KEY", raising=False)
    stato, _, corpo = _chiama("/api/analizza", "POST", {"url": "https://x.it/"})
    assert stato.startswith("500")
    assert "GOOGLE_API_KEY" in json.loads(corpo)["errore"]


def test_il_report_rifiuta_una_lista_vuota(monkeypatch):
    stato, _, corpo = _chiama("/api/report", "POST", {"pagine": []})
    assert stato.startswith("400")


def test_il_report_scarica_un_docx(monkeypatch, pagina):
    stato, intestazioni, corpo = _chiama("/api/report", "POST", {"pagine": [pagina]})
    assert stato.startswith("200")
    assert "wordprocessingml" in intestazioni["Content-Type"]
    assert intestazioni["Content-Disposition"].startswith("attachment")
    assert corpo[:2] == b"PK", "un .docx e' uno zip"


def test_lo_stato_dice_cosa_vede_il_server(monkeypatch):
    """Serve a distinguere 'non l'ho impostata' da 'non e' stata iniettata'."""
    monkeypatch.setenv("GOOGLE_API_KEY", "AIza-finta-1234")
    stato, _, corpo = _chiama("/api/stato")
    dati = json.loads(corpo)
    assert stato.startswith("200")
    assert dati["GOOGLE_API_KEY"] == "presente"


def test_lo_stato_non_dice_quanto_e_lunga_la_chiave(monkeypatch):
    """L'endpoint e' aperto come il resto dell'app: la lunghezza smaschererebbe un
    incollaggio troncato, ma la direbbe a chiunque abbia il link."""
    monkeypatch.setenv("GOOGLE_API_KEY", "AIza-finta-1234")
    assert "caratteri" not in _chiama("/api/stato")[2].decode("utf-8")


def test_lo_stato_non_espone_mai_il_valore_della_chiave(monkeypatch):
    monkeypatch.setenv("GOOGLE_API_KEY", "chiave-segretissima")
    for chiamata in (_chiama("/api/stato"), _chiama("/api/stato", "POST", {})):
        assert "segretissima" not in chiamata[2].decode("utf-8")


def test_lo_stato_riconosce_una_variabile_vuota(monkeypatch):
    monkeypatch.setenv("GOOGLE_API_KEY", "   ")
    assert "vuota" in json.loads(_chiama("/api/stato")[2])["GOOGLE_API_KEY"]


# --- protezione della quota -------------------------------------------------- #

def test_il_limite_di_richieste_scatta(monkeypatch):
    import app as applicazione
    monkeypatch.setattr(applicazione, "_conteggio", {})
    monkeypatch.setattr(applicazione, "LIMITE_RICHIESTE", 3)
    environ = {"REMOTE_ADDR": "1.2.3.4"}
    assert [applicazione._oltre_il_limite(environ) for _ in range(4)] ==         [False, False, False, True]


def test_il_limite_e_per_origine(monkeypatch):
    import app as applicazione
    monkeypatch.setattr(applicazione, "_conteggio", {})
    monkeypatch.setattr(applicazione, "LIMITE_RICHIESTE", 1)
    assert applicazione._oltre_il_limite({"REMOTE_ADDR": "1.1.1.1"}) is False
    assert applicazione._oltre_il_limite({"REMOTE_ADDR": "1.1.1.1"}) is True
    assert applicazione._oltre_il_limite({"REMOTE_ADDR": "2.2.2.2"}) is False


def test_la_finestra_scorre(monkeypatch):
    import app as applicazione
    monkeypatch.setattr(applicazione, "_conteggio", {})
    monkeypatch.setattr(applicazione, "LIMITE_RICHIESTE", 1)
    environ = {"REMOTE_ADDR": "3.3.3.3"}
    assert applicazione._oltre_il_limite(environ, adesso=0) is False
    assert applicazione._oltre_il_limite(environ, adesso=10) is True
    assert applicazione._oltre_il_limite(
        environ, adesso=applicazione.FINESTRA_SECONDI + 20) is False


def test_l_origine_viene_dal_proxy():
    import app as applicazione
    assert applicazione._origine({"HTTP_X_FORWARDED_FOR": "9.9.9.9, 10.0.0.1"}) == "9.9.9.9"
    assert applicazione._origine({"REMOTE_ADDR": "8.8.8.8"}) == "8.8.8.8"


def test_l_analisi_rifiuta_oltre_il_limite(monkeypatch):
    import app as applicazione
    monkeypatch.setenv("GOOGLE_API_KEY", "finta")
    monkeypatch.setattr(applicazione, "_conteggio", {})
    monkeypatch.setattr(applicazione, "LIMITE_RICHIESTE", 0)
    stato, _, corpo = _chiama("/api/analizza", "POST", {"url": "https://x.it/"})
    assert stato.startswith("429")
    assert "quota Google" in json.loads(corpo)["rimedio"]


# --- ripresa dopo una pagina fallita ----------------------------------------- #

def test_il_report_accetta_pagine_riuscite_e_fallite(monkeypatch, pagina, tmp_path):
    """Il browser manda anche le pagine mancate: il report deve dichiararle invece
    di presentare una lista piu' corta come se fosse completa."""
    fallita = {"template": "/news", "url": "https://www.bbc.com/news",
               "errore": "PageSpeed Insights: non e' riuscito ad analizzare la pagina."}
    stato, _, corpo = _chiama("/api/report", "POST", {"pagine": [pagina, fallita]})
    assert stato.startswith("200") and corpo[:2] == b"PK"

    percorso = tmp_path / "r.docx"
    percorso.write_bytes(corpo)
    doc = Document(str(percorso))
    testo = " ".join(p.text for p in doc.paragraphs)
    assert "Misurazione fallita" in testo, "il report deve dire quali pagine mancano"
    assert "bbc.com/news" in testo


def test_il_ciclo_dell_interfaccia_non_si_ferma_alla_prima_pagina_fallita():
    """Prima il ciclo faceva `break` e perdeva tutte le pagine successive.

    Il controllo e' sul sorgente perche' e' JavaScript e non c'e' un browser nei
    test: se qualcuno reintroduce l'uscita anticipata, questo test se ne accorge.
    """
    sorgente = (RADICE / "public" / "index.html").read_text(encoding="utf-8")
    inizio = sorgente.index("for (let i = 0; i < urls.length; i++)")
    fine = sorgente.index("$('avvia').disabled = false;")
    ciclo = sorgente[inizio:fine]
    assert "break" not in ciclo, "una pagina fallita non deve fermare le altre"
    assert "falliti.push" in ciclo, "la pagina fallita va accumulata, non persa"


def test_l_interfaccia_manda_al_report_anche_le_pagine_fallite():
    sorgente = (RADICE / "public" / "index.html").read_text(encoding="utf-8")
    assert "[...risultati, ...falliti]" in sorgente


def test_il_download_resta_possibile_con_pagine_fallite():
    """Il pulsante si disabilita solo se NON e' riuscita nessuna pagina."""
    sorgente = (RADICE / "public" / "index.html").read_text(encoding="utf-8")
    assert "$('scarica').disabled = risultati.length === 0;" in sorgente


# --- la lista degli interventi e' una sola per il sito ----------------------- #

def test_l_interfaccia_raggruppa_gli_interventi_per_tipo():
    """Su tre template, 10 interventi su 13 comparivano su tutti e tre: 20 schede
    su 37 erano ripetizioni del titolo. Il controllo e' sul sorgente perche' e'
    JavaScript e non c'e' un browser nei test."""
    sorgente = (RADICE / "public" / "index.html").read_text(encoding="utf-8")
    assert "function raggruppa(" in sorgente
    assert "function disegnaInterventi(" in sorgente
    assert "disegnaInterventi();" in sorgente, "va richiamata dopo ogni pagina"


def test_l_interfaccia_isola_i_bersagli_comuni_a_tutti_i_template():
    """I file presenti su ogni template sono il bundle condiviso: sistemarli una
    volta vale per tutto il sito, ed e' una informazione diversa dalle altre."""
    sorgente = (RADICE / "public" / "index.html").read_text(encoding="utf-8")
    assert "function elencoBersagli(" in sorgente
    assert "su tutti i ${gruppo.quanti} template" in sorgente


def test_la_scheda_di_pagina_non_ripete_piu_gli_interventi():
    sorgente = (RADICE / "public" / "index.html").read_text(encoding="utf-8")
    inizio = sorgente.index("function disegna(pagina)")
    fine = sorgente.index("function disegnaFallita(pagina)")
    scheda = sorgente[inizio:fine]
    assert "problemi(pagina.problemi" not in scheda,         "gli interventi stanno in una sezione sola, non ripetuti per pagina"


def test_url_valido_passa_la_validazione():
    assert valida_url("https://www.esempio.it/") is None
    assert valida_url("esempio.it") is not None


# --- l'attesa: la pagina deve dire che sta lavorando ------------------------- #

def _sorgente():
    return (RADICE / "public" / "index.html").read_text(encoding="utf-8")


def test_l_avanzamento_ha_una_riga_per_ogni_url():
    """Prima l'unico segno di vita era una riga di testo che cambiava una volta
    per URL: stava ferma fino a un minuto su un'analisi che ne richiede 20-60.
    Ora ogni URL ha il suo stato."""
    sorgente = _sorgente()
    assert "function disegnaAvanzamento(" in sorgente
    for stato in ("coda", "corso", "fatto", "errore"):
        assert f"{stato}:" in sorgente, f"manca lo stato '{stato}'"


def test_il_contatore_dei_secondi_scorre_durante_l_attesa():
    """Un numero che cambia ogni secondo e' cio' che distingue 'sta lavorando' da
    'si e' piantato'. Senza l'intervallo la riga resterebbe ferma."""
    sorgente = _sorgente()
    assert "setInterval(disegnaAvanzamento, 1000)" in sorgente
    assert "clearInterval(orologio)" in sorgente, "l'orologio va fermato a fine analisi"


def test_l_orologio_si_ferma_anche_quando_una_pagina_fallisce():
    """fermaOrologio() sta dopo il ciclo, non dentro il ramo riuscito: un errore
    sull'ultima pagina non deve lasciare un intervallo che gira per sempre."""
    sorgente = _sorgente()
    ciclo = sorgente[sorgente.index("for (let i = 0; i < urls.length; i++)"):]
    assert ciclo.index("fermaOrologio();") < ciclo.index("$('scarica').disabled")


# --- la sintesi in cima ------------------------------------------------------ #

def test_la_sintesi_conta_i_template_fuori_soglia_per_metrica():
    """Con cinque template la pagina supera i cinquemila pixel: senza una sintesi
    in cima bisogna scorrere tutto per sapere se qualcosa e' fuori soglia."""
    sorgente = _sorgente()
    assert "function disegnaVerdetto(" in sorgente
    assert "disegnaVerdetto();" in sorgente, "va richiamata dopo ogni pagina"
    assert "fuori soglia" in sorgente


def test_la_sintesi_non_usa_il_punteggio_psi():
    """ADR-001: il punteggio non entra in nessun giudizio. La sintesi giudica
    guardando le soglie dei Core Web Vitals, non lo score."""
    sorgente = _sorgente()
    inizio = sorgente.index("function disegnaVerdetto(")
    corpo = sorgente[inizio:sorgente.index("function disegnaAvanzamento(")]
    assert "punteggio" not in corpo and "score" not in corpo


def test_la_sintesi_indicizza_i_template():
    sorgente = _sorgente()
    assert 'href="#interventi"' in sorgente
    assert 'href="#template-' in sorgente


# --- il modulo si toglie di mezzo quando ci sono risultati ------------------- #

def test_il_modulo_si_comprime_a_fine_analisi():
    """Misurato: a risultati pronti il modulo occupava ancora tutto il primo
    schermo, e i dati cominciavano sotto la piega."""
    sorgente = _sorgente()
    assert "function comprimiModulo(" in sorgente
    assert "function espandiModulo(" in sorgente
    assert "comprimiModulo();" in sorgente


def test_una_pagina_fallita_si_puo_ritentare_da_sola():
    """Rifare tutte le pagine per una caduta costa altri minuti di attesa."""
    sorgente = _sorgente()
    assert "async function riprova(" in sorgente
    assert "riprova(" in sorgente[sorgente.index("function disegnaFallita("):]


def test_il_ritento_non_duplica_la_pagina_fra_i_falliti():
    """Se riesce, la pagina deve uscire dall'elenco degli errori: altrimenti il
    report la dichiarerebbe mancata pur avendone i dati."""
    sorgente = _sorgente()
    corpo = sorgente[sorgente.index("async function riprova("):
                     sorgente.index("async function scarica(")]
    assert "falliti.splice" in corpo


# --- memoria della pagina ----------------------------------------------------- #

def test_la_pagina_ricorda_gli_url_fra_una_sessione_e_l_altra():
    sorgente = _sorgente()
    assert "function ricorda(" in sorgente and "function ripristina(" in sorgente
    assert "localStorage" in sorgente


def test_la_pagina_non_chiede_nessuna_password():
    """Il tool non ha autenticazione: un campo password sarebbe un ostacolo che non
    protegge niente, e una password salvata nel browser sarebbe un segreto in piu'
    da custodire per zero beneficio."""
    sorgente = _sorgente().lower()
    for traccia in ("password", "sessionstorage"):
        assert traccia not in sorgente, f"la pagina nomina ancora '{traccia}'"


# --- il raggruppamento dei bersagli deve essere visibile --------------------- #

def test_il_raggruppamento_per_template_ha_una_regola_css():
    """Misurato: 'li.gruppo' non aveva nessuna regola, quindi l'intestazione del
    template era indistinguibile dalla riga di un file."""
    sorgente = _sorgente()
    assert ".gruppo" in sorgente[:sorgente.index("<body")], "la regola manca nel CSS"


# --- difetti visti negli screenshot, non dedotti ---------------------------- #

def test_hidden_vince_sulle_regole_con_display():
    """Visto a schermo: la barra compatta compariva a pagina vuota, con dentro
    'Rianalizza' e 'Scarica il report'. `[hidden]` e' display:none solo nel foglio
    del browser, e `.barra{display:flex}` lo scavalcava."""
    sorgente = _sorgente()
    assert "[hidden] { display:none !important; }" in sorgente


def test_il_modulo_si_chiude_all_avvio_non_alla_fine():
    """Misurato a 390px: con il modulo aperto l'elenco di avanzamento cominciava
    a 950px, cioe' fuori dal primo schermo per tutti i minuti in cui e' l'unica
    cosa che si muove. Le pagine misurate hanno richiesto 24, 66 e 89 secondi."""
    sorgente = _sorgente()
    corpo = sorgente[sorgente.index("async function avvia()"):]
    corpo = corpo[:corpo.index("async function riprova(")]
    assert corpo.index("comprimiModulo();") < corpo.index("avviaOrologio();")


def test_durante_l_analisi_non_si_puo_avviarne_una_seconda():
    """Il modulo e' chiuso ma i suoi comandi restano nella barra: senza questo,
    'Rianalizza' partirebbe sopra l'analisi in corso."""
    sorgente = _sorgente()
    assert "function comandiAttivi(" in sorgente
    assert "comandiAttivi(false);" in sorgente and "comandiAttivi(true);" in sorgente


def test_errore_e_consiglio_stanno_fuori_dal_modulo():
    """Chiuso il modulo sparirebbero anche loro, e il consiglio sui domini propri
    compare proprio a meta' analisi, dopo la prima pagina."""
    sorgente = _sorgente()
    modulo = sorgente[sorgente.index('<div class="riquadro" id="modulo">'):
                      sorgente.index('<!-- Fuori dal modulo')]
    assert 'id="errore"' not in modulo and 'id="consiglio"' not in modulo
    assert 'id="errore"' in sorgente and 'id="consiglio"' in sorgente
