"""
Test dei due renderer e delle convenzioni di naming.

Il run di prova non e' un file statico: viene composto qui dai fixture reali,
seguendo la stessa sequenza della CLI. Cosi' i test coprono anche la composizione,
e non c'e' un artefatto in piu' da tenere allineato a mano.
"""
import json
from pathlib import Path

import pytest
from docx import Document

from speed.cli import _nomi
from speed.config import Config, Template
from speed.core import consenso, diagnose, extract, thirdparty
from speed.io import crux, render, render_docx

FIXTURES = Path(__file__).resolve().parent.parent / "fixtures"
URL = "https://www.bbc.com/"
PROPRI = ["bbci.co.uk"]


def _carica(nome):
    return json.loads((FIXTURES / nome).read_text(encoding="utf-8"))


def _campo():
    """Voce `campo` come la compone la CLI, usando il parsing vero di crux.py."""
    record = crux.leggi_record(_carica("crux-bbc.json"))
    storico = crux.leggi_storico(_carica("crux-history-bbc.json"))
    return {"livello": "url", "metriche": record["metriche"],
            "storico": {"url": URL} | storico}


@pytest.fixture(scope="module")
def esecuzione():
    misurazioni = [extract.estrai(_carica(nome), URL, "PHONE", PROPRI)
                   for nome in ("psi-bbc-mobile-it.json", "psi-bbc-mobile-it-2.json")]
    accordo = consenso.combina(misurazioni)
    fatti = accordo.fatti
    campo = _campo()
    riepilogo = thirdparty.riepiloga(fatti.richieste, URL, PROPRI)
    problemi = diagnose.diagnostica(fatti, campo["metriche"], riepilogo, accordo)

    esecuzione = {
        "cliente": "Esempio", "sito": URL, "data": "2026-08-20", "form_factor": "PHONE",
        "pagine": [{
            "template": "Home", "url": URL, "fatti": fatti, "campo": campo,
            "terze_parti": riepilogo, "problemi": problemi,
            "misurazioni": accordo.ripetizioni, "concordi": accordo.concordi,
            "consenso": accordo.descrizione,
        }],
    }
    # I renderer lavorano sulla forma JSON, non sugli oggetti in memoria.
    from speed.cli import _serializza
    return json.loads(json.dumps(esecuzione, default=_serializza, ensure_ascii=False))


# --- HTML ------------------------------------------------------------------ #

def test_html_e_self_contained(esecuzione):
    h = render.html_report(esecuzione)
    assert "<script" not in h, "nessun JavaScript: il file deve aprirsi anche offline"
    assert "http://" not in h.split("<body>")[0], "nessuna risorsa esterna nell'head"
    assert "<svg" in h, "lo storico va reso come sparkline inline"


def test_html_dichiara_i_limiti(esecuzione):
    h = render.html_report(esecuzione)
    assert "28 giorni" in h, "va detto che il campo e' una media mobile"
    assert "non calibrata" in h or "non vengono usati" in h.replace("\n", " ")


def test_html_dichiara_la_provenienza_del_testo(esecuzione):
    h = render.html_report(esecuzione)
    assert "testo di Lighthouse" in h
    assert "classificazione su dati di campo" in h,         "i fixture CrUX hanno le fasi LCP: il report deve dire che vengono dal campo"


def test_html_riporta_le_risorse_colpevoli(esecuzione):
    h = render.html_report(esecuzione)
    assert "risorse" in h and "impatto" in h


def test_html_dichiara_quante_misurazioni(esecuzione):
    assert "misurazioni di laboratorio" in render.html_report(esecuzione)


# --- DOCX ------------------------------------------------------------------ #

def _testo_docx(percorso):
    doc = Document(str(percorso))
    parti = [p.text for p in doc.paragraphs]
    for tabella in doc.tables:
        for riga in tabella.rows:
            parti += [c.text for c in riga.cells]
    return "\n".join(parti)


def test_docx_si_genera(esecuzione, tmp_path):
    percorso = render_docx.docx_report(esecuzione, tmp_path / "r.docx")
    assert percorso.exists() and percorso.stat().st_size > 10_000


def test_docx_contiene_gli_stessi_fatti_dell_html(esecuzione, tmp_path):
    testo = _testo_docx(render_docx.docx_report(esecuzione, tmp_path / "r.docx"))
    for atteso in ("p75 reale", "testo di Lighthouse", "28 giorni",
                   "misurazioni di laboratorio", "Interventi"):
        assert atteso in testo, f"manca nel DOCX: {atteso}"


def test_docx_ha_la_tabella_del_campo(esecuzione, tmp_path):
    doc = Document(str(render_docx.docx_report(esecuzione, tmp_path / "r.docx")))
    assert doc.tables, "il campo va reso come tabella, non come testo"
    intestazione = [c.text for c in doc.tables[0].rows[0].cells]
    assert "Metrica" in intestazione and "p75 reale" in intestazione


def test_sparkline_a_blocchi():
    assert render_docx._sparkline([1, 2, 3, 4, 5])
    assert render_docx._sparkline([None, None]) == "", "serie vuota -> nessuna sparkline"
    assert render_docx._sparkline([5]) == "", "un punto solo non e' un andamento"


# --- naming ---------------------------------------------------------------- #

def _config(output):
    return Config(cliente="Cliente Uno", sito="https://x.it",
                  template=[Template(nome="Home", url="https://x.it/")], output=output)


def test_nome_senza_cliente_se_scrive_nella_cartella_cliente():
    _, report = _nomi(_config("C:/Clienti/Uno/05_Report"))
    assert "Cliente" not in report and report.startswith("Report velocità ")


def test_nome_con_cliente_se_scrive_in_out():
    _, report = _nomi(_config(""))
    assert "Cliente-Uno" in report


def test_nessuna_versione_nel_nome():
    for output in ("", "C:/Clienti/Uno"):
        dati, report = _nomi(_config(output))
        assert "_v" not in report and "finale" not in report.lower()
        assert dati != report, "dati e report non devono collidere"
