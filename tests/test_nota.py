"""
Test della nota tecnica per lo sviluppo.

E' il documento di consegna: pochi problemi accorpati per tema, in ordine, con la
citazione di PageSpeed per ognuno. Rispetto al riferimento completo (`render_md`)
la differenza non e' di formattazione ma di contenuto, e i test presidiano proprio
le tre cose che il modello decide da solo: quali audit stanno insieme, quanto vale
un tema, e quali frasi di sintesi sono autorizzate.
"""
import json
from pathlib import Path

import pytest
from docx import Document

from speed.cli import _stem_nota
from speed.core import nota
from speed.io import render_nota

RADICE = Path(__file__).resolve().parent.parent
FIXTURES = RADICE / "fixtures"


@pytest.fixture(scope="module")
def esecuzione():
    import sys
    sys.path.insert(0, str(RADICE / "scripts"))
    from esempi import esecuzione_di_prova
    return esecuzione_di_prova()


@pytest.fixture(scope="module")
def senza_campo(esecuzione):
    """Lo stesso run come si presenta su uno staging: nessun dato CrUX."""
    pagine = [{**p, "campo": {"livello": "assente", "metriche": {}}}
              if not p.get("errore") else p for p in esecuzione["pagine"]]
    return {**esecuzione, "pagine": pagine}


def _tema(temi, codice):
    return next(t for t in temi if t.codice == codice)


# --- accorpamento per tema ---------------------------------------------------- #

def test_gli_audit_dello_stesso_lavoro_finiscono_in_un_tema_solo(esecuzione):
    """unused-javascript, third-parties e la catena di richieste sono lo stesso
    lavoro per chi lo deve fare: una voce, non tre."""
    tema = _tema(nota.temi(esecuzione), "js-inutile")
    assert {"unused-javascript", "third-parties-insight",
            "network-dependency-tree-insight"} <= set(tema.audit)


def test_ogni_audit_finisce_in_un_tema_e_uno_solo(esecuzione):
    temi = nota.temi(esecuzione)
    tutti = [a for t in temi for a in t.audit]
    assert len(tutti) == len(set(tutti)), "un audit non puo' stare in due temi"


def test_il_numero_di_voci_e_leggibile_in_riunione(esecuzione):
    """E' tutta la ragione del documento: quaranta schede non si leggono."""
    assert len(nota.temi(esecuzione)) <= 10


def test_nessun_audit_sparisce_nell_accorpamento(esecuzione):
    """Ogni codice presente nei problemi dev'essere in un tema."""
    nei_temi = {a for t in nota.temi(esecuzione) for a in t.audit}
    nei_problemi = {p["codice"] for pagina in esecuzione["pagine"]
                    if not pagina.get("errore")
                    for p in pagina.get("problemi") or []}
    assert nei_problemi <= nei_temi


def test_il_tema_predefinito_raccoglie_cio_che_non_e_in_tabella():
    assert nota.tema_di("audit-mai-visto-prima") == nota.TEMA_PREDEFINITO
    assert nota.tema_di("bootup-time")[0] == "main-thread"
    assert nota.tema_di("lcp-resourceLoadDelay")[0] == "lcp"


# --- il valore del tema: massimo, non somma ----------------------------------- #

def test_il_peso_del_tema_e_il_massimo_non_la_somma(esecuzione):
    """bootup-time attribuisce agli script il lavoro che mainthread-work
    ripartisce per categoria: sommarli lo conterebbe due volte."""
    tema = _tema(nota.temi(esecuzione), "main-thread")
    massimi = []
    for pagina in esecuzione["pagine"]:
        if pagina.get("errore"):
            continue
        for problema in pagina["problemi"]:
            if problema["codice"] in tema.audit:
                massimi.append(nota._spreco(problema)[1])
    assert tema.ms_sprecati == pytest.approx(max(massimi))
    assert tema.ms_sprecati < sum(massimi), "sommare conterebbe due volte"


def test_il_titolo_dice_fino_a_perche_e_un_massimo(esecuzione):
    for tema in nota.temi(esecuzione):
        if tema.byte_sprecati >= 1024 * 100 or (
                tema.ms_sprecati >= 100 and tema.codice in nota.SUFFISSO_MS):
            assert "fino a" in tema.titolo, tema.codice


def test_i_millisecondi_senza_significato_dichiarato_non_si_mostrano(esecuzione):
    """In un tema misto i ms arrivano da audit diversi e non c'e' una frase vera
    che li descriva: dire "di risparmio stimato" sarebbe comodo e falso."""
    tema = _tema(nota.temi(esecuzione), "js-inutile")
    assert tema.ms_sprecati > 0, "il tema ha comunque dei millisecondi"
    assert "ms" not in tema.titolo and " s " not in tema.titolo


# --- gravita': una soglia dichiarata, non un giudizio ------------------------- #

def test_col_campo_disponibile_la_gravita_resta_quella_del_campo(esecuzione):
    """ADR-001: dichiarare BLOCCANTE un LCP che in lab e' 10 s e sugli utenti
    reali e' 1,1 s riporterebbe la priorita' al laboratorio di nascosto."""
    assert nota.quadro(esecuzione)["modalita"] == "campo"
    assert not [t for t in nota.temi(esecuzione) if t.gravita == "bloccante"]


def test_senza_campo_il_laboratorio_puo_alzare_a_bloccante(senza_campo):
    assert nota.quadro(senza_campo)["modalita"] == "laboratorio"
    bloccanti = [t.codice for t in nota.temi(senza_campo) if t.gravita == "bloccante"]
    assert "main-thread" in bloccanti and "lcp" in bloccanti


def test_i_temi_sono_ordinati_per_gravita(esecuzione):
    temi = nota.temi(esecuzione)
    ordine = [nota.ORDINE_GRAVITA[t.gravita] for t in temi]
    assert ordine == sorted(ordine)


# --- sintesi: regole dichiarate, non prosa ------------------------------------ #

def test_ogni_frase_di_sintesi_viene_da_una_regola(esecuzione):
    frasi = nota.sintesi(esecuzione)
    assert frasi, "sul run di prova almeno una regola deve scattare"
    assert len(frasi) <= len(nota.REGOLE_SINTESI)


def test_le_regole_non_scattano_senza_i_numeri():
    """Un run vuoto non deve produrre frasi: la condizione guarda i dati."""
    assert nota.sintesi({"pagine": []}) == []


def test_la_frase_sul_cls_richiede_che_il_cls_sia_davvero_a_posto(esecuzione):
    frasi = " ".join(nota.sintesi(esecuzione))
    m = nota.misure(esecuzione)
    if "non si perde nel layout" in frasi:
        assert m.cls_entro_soglia == m.pagine


def test_i_numeri_delle_frasi_vengono_dalle_misure(esecuzione):
    m = nota.misure(esecuzione)
    frasi = " ".join(nota.sintesi(esecuzione))
    if m.peggiore:
        assert m.peggiore in frasi


# --- quadro di sintesi -------------------------------------------------------- #

def test_il_quadro_porta_le_metriche_di_laboratorio(esecuzione):
    q = nota.quadro(esecuzione)
    for sigla in ("FCP", "LCP", "TBT", "TTI", "CLS"):
        assert sigla in q["intestazioni"]
    assert len(q["righe"]) == 2


def test_col_campo_il_quadro_aggiunge_le_colonne_reali(esecuzione, senza_campo):
    assert "LCP campo" in nota.quadro(esecuzione)["intestazioni"]
    assert "LCP campo" not in nota.quadro(senza_campo)["intestazioni"]


def test_le_pagine_fallite_non_entrano_nel_quadro(esecuzione):
    q = nota.quadro(esecuzione)
    assert q["pagine"] == 2, "la pagina fallita non e' una riga di metriche"


# --- il documento Word -------------------------------------------------------- #

def _testo(percorso):
    documento = Document(str(percorso))
    parti = [p.text for p in documento.paragraphs]
    for tabella in documento.tables:
        for riga in tabella.rows:
            parti.append(" | ".join(c.text for c in riga.cells))
    return "\n".join(parti)


def test_il_word_si_genera(esecuzione, tmp_path):
    percorso = render_nota.nota_docx(esecuzione, tmp_path / "nota.docx")
    assert percorso.exists() and percorso.stat().st_size > 10_000


def test_il_word_ha_le_sezioni_del_documento_di_riferimento(esecuzione, tmp_path):
    testo = _testo(render_nota.nota_docx(esecuzione, tmp_path / "n.docx"))
    for sezione in ("Performance sito", "nota tecnica per lo sviluppo",
                    "Quadro di sintesi", "I problemi, in ordine di priorita'",
                    "Ordine di lavorazione", "Nota di lettura"):
        assert sezione in testo, sezione


def test_ogni_tema_cita_psi_e_rimanda_alla_documentazione(esecuzione, tmp_path):
    testo = _testo(render_nota.nota_docx(esecuzione, tmp_path / "n.docx"))
    assert testo.count("PSI — ") >= 8
    assert testo.count("Documentazione Google:") >= 8
    for url in ("developer.chrome.com", "https://"):
        assert url in testo


def test_i_link_sono_quelli_di_lighthouse(esecuzione, tmp_path):
    """Nessun URL che non venga dai dati del run."""
    import re
    testo = _testo(render_nota.nota_docx(esecuzione, tmp_path / "n.docx"))
    leciti = {o.get("documentazione") for pagina in esecuzione["pagine"]
              for o in (pagina.get("fatti") or {}).get("opportunita") or []
              if o.get("documentazione")}
    for url in re.findall(r"Documentazione Google:\s+(\S+)", testo):
        assert url in leciti, url


def test_la_provenienza_del_testo_e_dichiarata_in_testa(esecuzione, tmp_path):
    """ADR-004: chi legge deve sapere quale riga viene da Google e quale da noi."""
    testo = _testo(render_nota.nota_docx(esecuzione, tmp_path / "n.docx"))
    assert "restituite da PageSpeed Insights, citate senza integrazioni" in testo
    assert "sono invece nostri" in testo


def test_il_documento_dichiara_la_modalita(esecuzione, senza_campo, tmp_path):
    con = _testo(render_nota.nota_docx(esecuzione, tmp_path / "a.docx"))
    assert "poggia sui dati di campo CrUX" in con
    senza = _testo(render_nota.nota_docx(senza_campo, tmp_path / "b.docx"))
    assert "CrUX non ha dati di campo" in senza
    assert "poggia sul laboratorio" in senza


def test_il_punteggio_psi_non_compare(esecuzione, tmp_path):
    """ADR-001, e qui non c'e' nemmeno il piede di vetrina: e' una nota di lavoro."""
    testo = _testo(render_nota.nota_docx(esecuzione, tmp_path / "n.docx"))
    assert "non entra in nessuna delle valutazioni" in testo
    for pagina in esecuzione["pagine"]:
        score = (pagina.get("fatti") or {}).get("performance_score")
        if score is not None:
            assert f"punteggio {score}" not in testo.lower()


def test_la_pagina_fallita_e_dichiarata(esecuzione, tmp_path):
    testo = _testo(render_nota.nota_docx(esecuzione, tmp_path / "n.docx"))
    assert "Non misurate" in testo and "Video" in testo


def test_gli_audit_non_assegnati_portano_il_motivo(esecuzione, tmp_path):
    testo = _testo(render_nota.nota_docx(esecuzione, tmp_path / "n.docx"))
    assert "Fuori dal master plan" in testo
    assert "artefatto di dati" in testo


def test_il_documento_nomina_le_chiavi_degli_audit(esecuzione, tmp_path):
    """Servono a rilanciare Lighthouse: restano anche nella versione concisa."""
    testo = _testo(render_nota.nota_docx(esecuzione, tmp_path / "n.docx"))
    assert "Audit PageSpeed:" in testo
    assert "cache-insight" in testo and "bootup-time" in testo


def test_un_run_senza_problemi_non_esplode(tmp_path):
    vuoto = {"cliente": "X", "sito": "https://x.it", "data": "2026-01-01",
             "form_factor": "PHONE", "pagine": []}
    testo = _testo(render_nota.nota_docx(vuoto, tmp_path / "n.docx"))
    assert "Nessun audit oltre soglia" in testo


# --- ADR-004: nessuna raccomandazione scritta da noi -------------------------- #

def test_nessuna_azione_scritta_a_mano(esecuzione, tmp_path):
    testo = _testo(render_nota.nota_docx(esecuzione, tmp_path / "n.docx")).lower()
    for inventata in ("inventario dei tag", "valutare una facade", "digital pr",
                      "consigliamo", "suggeriamo", "dovresti", "bisogna"):
        assert inventata not in testo


def test_le_citazioni_psi_sono_testo_di_lighthouse(esecuzione, tmp_path):
    testo = _testo(render_nota.nota_docx(esecuzione, tmp_path / "n.docx"))
    descrizioni = {o["descrizione"] for pagina in esecuzione["pagine"]
                   for o in (pagina.get("fatti") or {}).get("opportunita") or []
                   if o.get("descrizione")}
    citate = [d for d in descrizioni if d in testo]
    assert len(citate) >= 8, "le citazioni devono essere verbatim"


def test_i_modelli_di_titolo_non_contengono_istruzioni():
    """Il titolo e' nostro ma non e' un'istruzione: dice cosa non va, non cosa fare."""
    import re
    sorgente = (RADICE / "speed" / "core" / "nota.py").read_text(encoding="utf-8")
    letterali = " ".join(re.findall(r'"([^"]{4,})"', sorgente)).lower()
    for imperativo in ("riduci ", "elimina ", "rimuovi ", "ottimizza ",
                       "comprimi ", "sostituisci "):
        assert imperativo not in letterali, imperativo


# --- naming -------------------------------------------------------------------- #

def test_il_nome_del_file_e_quello_con_cui_circolano():
    assert _stem_nota("Report velocità 20082026") == "Interventi Performance 20082026"
    assert _stem_nota("Report velocità Ferroli 20082026") == \
        "Interventi Performance Ferroli 20082026"


# --- il quadro non ripete l'URL ---------------------------------------------- #

@pytest.fixture(scope="module")
def dal_browser(esecuzione):
    """Lo stesso run come arriva dalla versione web: nessun nome di template."""
    pagine = [{**p, "template": p.get("url", "")} for p in esecuzione["pagine"]]
    return {**esecuzione, "pagine": pagine}


def test_con_i_nomi_la_colonna_template_resta(esecuzione):
    assert nota.quadro(esecuzione)["intestazioni"][:2] == ["Template", "URL"]


def test_senza_nomi_la_colonna_template_sparisce(dal_browser):
    """Il sintomo: Template e URL contenevano la stessa stringa su ogni riga."""
    quadro = nota.quadro(dal_browser)
    assert quadro["intestazioni"][0] == "URL"
    assert "Template" not in quadro["intestazioni"]
    for riga in quadro["righe"]:
        assert riga[0].startswith("http"), "la prima colonna e' l'URL"
        assert riga.count(riga[0]) == 1, "l'URL non si ripete nella riga"


def test_le_intestazioni_e_le_righe_restano_allineate(esecuzione, dal_browser):
    for run in (esecuzione, dal_browser):
        quadro = nota.quadro(run)
        for riga in quadro["righe"]:
            assert len(riga) == len(quadro["intestazioni"])


def test_i_titoli_dei_temi_non_gridano_l_url_intero(dal_browser):
    """Senza nome il tema si intestava con l'URL completo in maiuscolo."""
    for tema in nota.temi(dal_browser):
        for nome in tema.template:
            assert not nome.startswith("http"), nome


# --- una raccomandazione, una riga -------------------------------------------- #

def test_due_audit_con_la_stessa_frase_si_leggono_una_volta():
    """Il sintomo: nel tema LCP comparivano due righe consecutive con la stessa
    identica citazione, "Deve essere applicata fetchpriority=high"."""
    frase = "Deve essere applicata fetchpriority=high"
    accorpate = nota.accorpa_citazioni([
        ("Il browser scopre la risorsa LCP tardi", frase, "https://g.dev/a"),
        ("Il tempo se ne va nel download della risorsa LCP", frase, ""),
    ])
    assert len(accorpate) == 1
    titolo, testo, url = accorpate[0]
    assert testo == frase
    assert "scopre la risorsa LCP tardi" in titolo
    assert "download della risorsa LCP" in titolo
    assert url == "https://g.dev/a", "il primo link utile non si perde"


def test_la_riserva_sopravvive_all_accorpamento():
    """La riserva riguarda una delle due misure e va detta lo stesso."""
    frase = "Deve essere applicata fetchpriority=high"
    accorpate = nota.accorpa_citazioni([
        ("Il browser scopre la risorsa LCP tardi", frase, ""),
        ("Il tempo se ne va nel download [misurazioni discordanti]", frase, ""),
    ])
    assert "[misurazioni discordanti]" in accorpate[0][0]


def test_citazioni_diverse_restano_distinte():
    accorpate = nota.accorpa_citazioni([
        ("Riduci il CSS", "Riduci le regole inutilizzate.", "https://g.dev/css"),
        ("Minimizza CSS", "Minimizza i file CSS.", "https://g.dev/min"),
    ])
    assert len(accorpate) == 2


def test_nessun_tema_ripete_una_citazione(esecuzione):
    for tema in nota.temi(esecuzione):
        testi = [testo for _titolo, testo, _url in tema.citazioni]
        assert len(testi) == len(set(testi)), tema.codice


# --- «oltre soglia» conta le pagine, sul campo -------------------------------- #

def _con_campo(esecuzione, valori):
    """Lo stesso run con l'LCP di campo forzato pagina per pagina.

    `valori` e' una lista lunga quanto le pagine riuscite: un numero, oppure None
    per togliere del tutto i dati di campo a quella pagina.
    """
    pagine, indice = [], 0
    for pagina in esecuzione["pagine"]:
        if pagina.get("errore"):
            pagine.append(pagina)
            continue
        valore = valori[indice]
        indice += 1
        if valore is None:
            pagine.append({**pagina, "campo": {"livello": "assente", "metriche": {}}})
        else:
            metriche = dict((pagina.get("campo") or {}).get("metriche") or {})
            metriche["largest_contentful_paint"] = valore
            pagine.append({**pagina, "campo": {**(pagina.get("campo") or {}),
                                               "metriche": metriche}})
    return {**esecuzione, "pagine": pagine}


def test_oltre_soglia_si_conta_sul_campo_non_sugli_audit(esecuzione):
    """Il sintomo: "LCP oltre soglia su 2 template" mentre la tabella CrUX dello
    stesso documento dava LCP buono su tutte e due."""
    conto = nota.conta_oltre_soglia(esecuzione, "LCP")
    assert conto.con_campo == 2 and conto.oltre_campo == 0
    lcp = _tema(nota.temi(esecuzione), "lcp")
    assert "oltre soglia" not in lcp.titolo, lcp.titolo
    assert "audit non superat" in lcp.titolo


def test_con_il_campo_oltre_soglia_il_titolo_lo_dice(esecuzione):
    run = _con_campo(esecuzione, [3684, 922])
    conto = nota.conta_oltre_soglia(run, "LCP")
    assert (conto.oltre_campo, conto.con_campo, conto.senza_campo) == (1, 2, 0)
    lcp = _tema(nota.temi(run), "lcp")
    assert lcp.titolo == "LCP oltre soglia sul campo su 1 template di 2"


def test_le_pagine_senza_campo_si_contano_a_parte(esecuzione):
    """Il caso Pluxee: due pagine con CrUX, il blog senza."""
    run = _con_campo(esecuzione, [3684, None])
    conto = nota.conta_oltre_soglia(run, "LCP")
    assert (conto.oltre_campo, conto.con_campo, conto.senza_campo) == (1, 1, 1)
    lcp = _tema(nota.temi(run), "lcp")
    assert "1 senza dati di campo" in lcp.titolo, lcp.titolo


def test_senza_campo_ovunque_il_titolo_dichiara_il_laboratorio(senza_campo):
    lcp = _tema(nota.temi(senza_campo), "lcp")
    assert "in laboratorio" in lcp.titolo, lcp.titolo


def test_il_proxy_non_diventa_un_affermazione_sulla_metrica():
    """TBT e' mappato su INP come proxy per calibrare la priorita'. Non basta a
    dire "TBT oltre soglia": e' un'altra metrica."""
    assert "TBT" not in nota.CAMPO_DELLA_SIGLA
    assert nota.CAMPO_DELLA_SIGLA["LCP"] == "largest_contentful_paint"


# --- la gravita' distingue ---------------------------------------------------- #

def test_la_gravita_distingue_gli_interventi(esecuzione):
    """Il sintomo: su una scansione reale otto voci su nove erano MEDIA, e la
    colonna non diceva da dove cominciare."""
    gravita = [t.gravita for t in nota.temi(esecuzione)]
    assert len(set(gravita)) >= 3, gravita
    assert gravita == sorted(gravita, key=lambda g: nota.ORDINE_GRAVITA[g]), \
        "la tabella resta ordinata dalla piu' grave"


def test_l_intervento_piu_grosso_sta_in_cima(esecuzione):
    """3 MB di JavaScript inutilizzato non possono stare alla pari con 126 ms
    di reflow."""
    temi = nota.temi(esecuzione)
    assert temi[0].codice == "js-inutile"
    reflow = _tema(temi, "reflow")
    assert nota.ORDINE_GRAVITA[temi[0].gravita] < nota.ORDINE_GRAVITA[reflow.gravita]


def test_la_taglia_viene_da_soglie_dichiarate():
    grosso = nota.Tema(codice="x", titolo="", gravita="media",
                       byte_sprecati=nota.BYTE_GROSSO)
    medio = nota.Tema(codice="x", titolo="", gravita="media",
                      byte_sprecati=nota.BYTE_MEDIO)
    piccolo = nota.Tema(codice="x", titolo="", gravita="media",
                        byte_sprecati=nota.BYTE_MEDIO - 1)
    assert nota.taglia_di(grosso) == "grossa"
    assert nota.taglia_di(medio) == "media"
    assert nota.taglia_di(piccolo) == "piccola"
    lento = nota.Tema(codice="x", titolo="", gravita="media",
                      ms_sprecati=nota.MS_GROSSO)
    assert nota.taglia_di(lento) == "grossa"


def test_un_tema_con_metrica_si_misura_sulla_metrica():
    """L'LCP non si pesa in kilobyte: la taglia e' quante pagine sono oltre
    soglia, non quanti byte dichiara l'audit."""
    tutte = nota.Tema(codice="lcp", titolo="", gravita="media",
                      soglia=nota.ContoSoglia(oltre_campo=3, con_campo=3))
    alcune = nota.Tema(codice="lcp", titolo="", gravita="media",
                       soglia=nota.ContoSoglia(oltre_campo=1, con_campo=3))
    nessuna = nota.Tema(codice="lcp", titolo="", gravita="media",
                        soglia=nota.ContoSoglia(oltre_campo=0, con_campo=3))
    assert nota.taglia_di(tutte) == "grossa"
    assert nota.taglia_di(alcune) == "media"
    assert nota.taglia_di(nessuna) == "piccola"


def test_col_campo_buono_nessun_tema_supera_media(esecuzione):
    """ADR-001: il campo tiene il tetto. Un tema grosso in laboratorio non puo'
    diventare ALTA se gli utenti reali stanno bene."""
    for urgenza, per_taglia in nota.GRAVITA_COMBINATA.items():
        if urgenza != "bassa":
            continue
        for gravita in per_taglia.values():
            assert nota.ORDINE_GRAVITA[gravita] >= nota.ORDINE_GRAVITA["media"]


def test_bloccante_non_si_combina(senza_campo):
    """BLOCCANTE ha una soglia sua ed e' gia' il massimo: la taglia non lo
    abbassa."""
    temi = nota.temi(senza_campo)
    bloccanti = [t for t in temi if t.urgenza == "bloccante"]
    assert bloccanti, "il caso ha senso solo se qualcuno scatta"
    for tema in bloccanti:
        assert tema.gravita == "bloccante"


def test_la_gravita_resta_una_regola_non_un_giudizio():
    """ADR-004: chi legge deve poter rifare il conto. Le due meta' sono
    entrambe enumerate."""
    assert set(nota.GRAVITA_COMBINATA) == {"alta", "media", "bassa"}
    for per_taglia in nota.GRAVITA_COMBINATA.values():
        assert set(per_taglia) == {"grossa", "media", "piccola"}

