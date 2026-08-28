"""
Nota tecnica per lo sviluppo, in Word.

E' il documento che si consegna: poche pagine, i problemi accorpati per tema e
ordinati, con la citazione di PageSpeed e il rimando alla documentazione Google
per ognuno. Il contenuto lo decide `core/nota.py`; qui si impagina soltanto.

Perche' Word e non Markdown: e' il formato in cui questi documenti circolano fra
noi e il cliente, ed e' quello in cui vengono annotati. Il Markdown resta per il
documento di riferimento completo (`render_md`), che ha un altro uso — il ticket
e la PR — e un altro contenuto: li' c'e' tutto, qui c'e' cio' su cui si lavora.

Lo stile e' quello di `render_docx`: stesse dimensioni, stessi colori, stesse
tabelle. Sono due documenti dello stesso studio.
"""
from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from ..core import nota
from ..core.aggregazione import etichetta_pagina
from ..core.soglie import ETICHETTE, formatta, giudizio

GRIGIO = RGBColor(0x6B, 0x72, 0x80)
NERO = RGBColor(0x1F, 0x23, 0x28)
COLORE_GRAVITA = {
    "bloccante": RGBColor(0x8B, 0x11, 0x0B),
    "alta": RGBColor(0xB4, 0x23, 0x18),
    "media": RGBColor(0xA1, 0x62, 0x07),
    "bassa": GRIGIO,
}

# Le due frasi che dichiarano la provenienza del testo. Stanno in testa perche'
# reggono tutto il documento: ADR-004 vuole che chi legge sappia, senza doverlo
# chiedere, quale riga viene da Google e quale da noi.
PROVENIENZA = (
    "Le indicazioni di intervento riportate in questo documento sono quelle "
    "restituite da PageSpeed Insights, citate senza integrazioni, con il rimando "
    "alla documentazione Google corrispondente. La diagnosi tecnica e la scelta "
    "della soluzione restano allo sviluppo."
)
NOSTRO = (
    "Titoli, accorpamento per tema, gravita' e frasi di sintesi sono invece "
    "nostri: derivano per regola dai numeri misurati, non da una valutazione caso "
    "per caso."
)

NOTA_LABORATORIO = (
    "Il test e' stato eseguito in laboratorio, con dati simulati: rete e CPU "
    "rallentate, non traffico reale. Per queste pagine CrUX non ha dati di campo "
    "— e' quello che succede su un ambiente non pubblico o senza traffico "
    "sufficiente — quindi l'ordine di priorita' qui poggia sul laboratorio. Prima "
    "e dopo gli interventi vale la pena confrontare i dati di campo, che sono "
    "quelli che Google usa per il ranking."
)
# Quando il campo c'e' su alcune pagine e non su altre. Senza questa riga il
# documento dichiara che l'ordine poggia sul campo e poi mostra "n/d" nel quadro,
# lasciando al lettore il compito di capire cosa significhi per quelle pagine.
SENZA_CAMPO = (
    "Senza dati di campo: {pagine}. CrUX non li espone quando il traffico non "
    "basta, ed e' il caso delle pagine nuove o poco visitate. Per queste la "
    "priorita' degli interventi non e' calibrata sugli utenti reali: vale il "
    "valore predefinito, «media». Nei conteggi «oltre soglia» si contano a "
    "parte, sul laboratorio, che e' un'altra misura e non si somma alla prima."
)
NOTA_CAMPO = (
    "I valori di laboratorio servono a capire dove si perde il tempo; l'ordine di "
    "priorita' poggia sui dati di campo CrUX, che sono gli utenti reali su 28 "
    "giorni ed e' cio' che Google usa per il ranking. Le due misure non "
    "coincidono, ed e' normale: il laboratorio simula un dispositivo lento su "
    "rete lenta."
)


def _p(doc, testo="", size=10, colore=NERO, corsivo=False, grassetto=False,
       spazio_dopo=4, spazio_prima=0):
    par = doc.add_paragraph()
    run = par.add_run(testo)
    run.font.size = Pt(size)
    run.font.color.rgb = colore
    run.italic = corsivo
    run.bold = grassetto
    par.paragraph_format.space_after = Pt(spazio_dopo)
    par.paragraph_format.space_before = Pt(spazio_prima)
    return par


def _etichettato(doc, etichetta: str, testo: str, size=9.5):
    """Un paragrafo "Evidenza: ...", con l'etichetta in grassetto.

    E' la forma del documento di riferimento: l'occhio trova la parola e salta al
    dato, senza leggere il paragrafo intero.
    """
    par = doc.add_paragraph()
    run = par.add_run(f"{etichetta}:  ")
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = NERO
    run = par.add_run(testo)
    run.font.size = Pt(size)
    run.font.color.rgb = NERO
    par.paragraph_format.space_after = Pt(3)
    return par


def _tabella(doc, intestazioni: list, righe: list, larghezze=None):
    tabella = doc.add_table(rows=1, cols=len(intestazioni))
    tabella.style = "Table Grid"
    for cella, testo in zip(tabella.rows[0].cells, intestazioni):
        run = cella.paragraphs[0].add_run(str(testo))
        run.bold = True
        run.font.size = Pt(8)
    for riga in righe:
        celle = tabella.add_row().cells
        for indice, testo in enumerate(riga):
            run = celle[indice].paragraphs[0].add_run(str(testo))
            run.font.size = Pt(8)
            if indice > 1:
                celle[indice].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tabella


# --------------------------------------------------------------------------- #
#  Sezioni
# --------------------------------------------------------------------------- #

def _testata(doc, esecuzione: dict):
    cliente = esecuzione.get("cliente", "")
    doc.add_heading(f"{cliente} — Performance sito", level=0)
    _p(doc, "Interventi da mettere a piano — nota tecnica per lo sviluppo",
       size=11, colore=GRIGIO, spazio_dopo=12)

    pagine = [p for p in esecuzione.get("pagine") or [] if not p.get("errore")]
    primo = (pagine[0].get("fatti") or {}) if pagine else {}
    strategia = "mobile" if esecuzione.get("form_factor") != "DESKTOP" else "desktop"
    versione = primo.get("lighthouse_version", "")

    _p(doc, f"Test PageSpeed Insights / Lighthouse {versione} del "
            f"{esecuzione.get('data', '')} su {esecuzione.get('sito', '')}, "
            f"{len(pagine)} template di pagina, profilo {strategia} "
            f"(throttling di rete simulato, CPU rallentata). I valori di "
            f"laboratorio sono quelli scalati dalla simulazione: servono a "
            f"confrontare le pagine fra loro, non come tempi reali su dispositivo.",
       size=9.5, spazio_dopo=6)
    _p(doc, PROVENIENZA, size=9.5, spazio_dopo=4)
    _p(doc, NOSTRO, size=9.5, colore=GRIGIO, spazio_dopo=12)


def _quadro(doc, quadro: dict, esecuzione: dict):
    doc.add_heading("Quadro di sintesi", level=1)
    _tabella(doc, quadro["intestazioni"], quadro["righe"])

    for frase in nota.sintesi(esecuzione):
        _p(doc, frase, size=9.5, spazio_dopo=4)

    fallite = [p for p in esecuzione.get("pagine") or [] if p.get("errore")]
    if fallite:
        _p(doc, "Non misurate: " + ", ".join(
            f"{etichetta_pagina(p)} ({p.get('errore', '')[:80]})" for p in fallite),
           size=9, colore=GRIGIO, corsivo=True, spazio_dopo=8)


def _tema(doc, indice: int, tema, esecuzione: dict):
    titolo = doc.add_heading(f"{indice:02d}  {tema.titolo}", level=2)
    for run in titolo.runs:
        run.font.color.rgb = COLORE_GRAVITA.get(tema.gravita, NERO)

    pagine = ("TUTTI I TEMPLATE" if len(tema.template) == tema.totale_template
              else " E ".join(t.upper() for t in tema.template))
    responsabili = ", ".join(tema.responsabili)
    _p(doc, f"{nota.ETICHETTA_GRAVITA.get(tema.gravita, '')} · {pagine}"
            + (f" · interviene: {responsabili}" if responsabili else ""),
       size=8.5, grassetto=True, colore=COLORE_GRAVITA.get(tema.gravita, GRIGIO),
       spazio_dopo=6)

    _etichettato(doc, "Evidenza", nota.evidenza_di_apertura(tema))

    # Le evidenze sono dati misurati, gia' raggruppati per etichetta dal modello:
    # prima cio' che vale su tutti i template, poi cio' che e' della singola pagina.
    for etichetta, testo in tema.evidenze:
        _etichettato(doc, etichetta, testo)

    for titolo_audit, testo, url in tema.citazioni:
        _etichettato(doc, f"PSI — {titolo_audit}", f"“{testo}”")
        if url:
            _p(doc, f"Documentazione Google:  {url}", size=8.5, colore=GRIGIO,
               spazio_dopo=3)

    for titolo_audit, motivo in tema.esclusioni:
        _p(doc, f"Fuori dal master plan:  {titolo_audit} — {motivo}",
           size=8.5, colore=GRIGIO, corsivo=True, spazio_dopo=3)

    _p(doc, "Audit PageSpeed: " + ", ".join(tema.audit),
       size=8, colore=GRIGIO, spazio_dopo=10)


def _ordine(doc, temi_ordinati: list):
    doc.add_heading("Ordine di lavorazione", level=1)
    _p(doc, "La gravita' combina due cose misurate: cosa dicono di quella "
            "metrica gli utenti reali su CrUX, e quanto pesa l'intervento "
            "secondo i numeri di Lighthouse. Con il campo «buono» un tema non "
            "supera MEDIA, per quanto grosso sia in laboratorio. «Fino a» "
            "perche' il valore e' il massimo fra gli audit del tema e fra le "
            "pagine, non una somma.", size=9, colore=GRIGIO, spazio_dopo=6)
    _tabella(doc, ["#", "Intervento", "Pagine", "Gravita'", "Peso del problema"],
             nota.ordine_di_lavorazione(temi_ordinati))


def _nota_di_lettura(doc, quadro: dict, esecuzione: dict):
    doc.add_heading("Nota di lettura", level=1)
    _p(doc, NOTA_LABORATORIO if quadro["modalita"] == "laboratorio" else NOTA_CAMPO,
       size=9.5, spazio_dopo=6)

    if quadro["modalita"] == "campo" and quadro["senza_campo"]:
        _p(doc, SENZA_CAMPO.format(pagine=", ".join(quadro["senza_campo"])),
           size=9.5, spazio_dopo=6)

    if quadro["modalita"] == "campo":
        righe = []
        for pagina in [p for p in esecuzione.get("pagine") or [] if not p.get("errore")]:
            campo = (pagina.get("campo") or {}).get("metriche") or {}
            for metrica in ("largest_contentful_paint", "interaction_to_next_paint",
                            "cumulative_layout_shift"):
                valore = campo.get(metrica)
                if valore is not None:
                    righe.append([etichetta_pagina(pagina), ETICHETTE[metrica],
                                  formatta(metrica, valore),
                                  giudizio(metrica, valore).replace("_", " ")])
        if righe:
            _p(doc, "Dati di campo CrUX, p75 su 28 giorni:", size=9,
               grassetto=True, spazio_dopo=4)
            _tabella(doc, ["Template", "Metrica", "p75", "Giudizio"], righe)

    _p(doc, "Il punteggio PageSpeed non compare in questo documento: e' rumoroso, "
            "cambia fra due chiamate identiche, e non entra in nessuna delle "
            "valutazioni qui sopra.", size=9, colore=GRIGIO, spazio_dopo=4)


# --------------------------------------------------------------------------- #
#  Composizione
# --------------------------------------------------------------------------- #

def nota_docx(esecuzione: dict, percorso):
    doc = Document()
    for stile in ("Normal",):
        doc.styles[stile].font.name = "Calibri"
        doc.styles[stile].font.size = Pt(10)

    quadro = nota.quadro(esecuzione)
    temi_ordinati = nota.temi(esecuzione)

    _testata(doc, esecuzione)
    _quadro(doc, quadro, esecuzione)

    doc.add_heading("I problemi, in ordine di priorita'", level=1)
    if not temi_ordinati:
        _p(doc, "Nessun audit oltre soglia sulle pagine misurate.", size=10)
    for indice, tema in enumerate(temi_ordinati, start=1):
        _tema(doc, indice, tema, esecuzione)

    if temi_ordinati:
        _ordine(doc, temi_ordinati)
    _nota_di_lettura(doc, quadro, esecuzione)

    doc.save(str(percorso))
    return percorso
