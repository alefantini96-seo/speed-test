"""
Report in formato Word.

Stesso contenuto dell'HTML: cambia il contenitore, non l'analisi. Entrambi i
renderer partono dalla forma JSON del run, quindi restano allineati per costruzione.

Le sparkline dello storico sono rese con caratteri di blocco invece che con
un'immagine: nessuna dipendenza grafica, e il documento resta modificabile.
"""
from __future__ import annotations

from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, RGBColor

from ..core.extract import FASI_IT
from ..core.soglie import ETICHETTE, SOGLIE, formatta, giudizio
from ..core.thirdparty import etichetta_tipo

GRIGIO = RGBColor(0x6B, 0x72, 0x80)
NERO = RGBColor(0x1F, 0x23, 0x28)
COLORI = {
    "buono": RGBColor(0x1A, 0x7F, 0x4B),
    "da_migliorare": RGBColor(0xA1, 0x62, 0x07),
    "scarso": RGBColor(0xB4, 0x23, 0x18),
    "sconosciuto": GRIGIO,
}
PAROLA = {"buono": "buono", "da_migliorare": "da migliorare", "scarso": "scarso",
          "sconosciuto": "n/d"}
GRAVITA = {"alta": COLORI["scarso"], "media": COLORI["da_migliorare"], "bassa": GRIGIO}

BLOCCHI = "▁▂▃▄▅▆▇█"


def _sparkline(valori: list, punti: int = 20) -> str:
    """Serie -> caratteri di blocco. Scala su min/max della serie stessa."""
    puliti = [v for v in valori if v is not None]
    if len(puliti) < 2:
        return ""
    passo = max(1, len(valori) // punti)
    campione = [v for v in valori[::passo] if v is not None]
    lo, hi = min(puliti), max(puliti)
    span = (hi - lo) or 1
    return "".join(BLOCCHI[min(int((v - lo) / span * (len(BLOCCHI) - 1)), len(BLOCCHI) - 1)]
                   for v in campione)


def _p(doc, testo="", size=10, colore=NERO, corsivo=False, grassetto=False, spazio_dopo=4):
    par = doc.add_paragraph()
    run = par.add_run(testo)
    run.font.size = Pt(size)
    run.font.color.rgb = colore
    run.italic = corsivo
    run.bold = grassetto
    par.paragraph_format.space_after = Pt(spazio_dopo)
    return par


def _avviso(doc, titolo, testo):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(8)
    par.paragraph_format.space_after = Pt(10)
    if titolo:
        run = par.add_run(f"{titolo} ")
        run.bold = True
        run.font.size = Pt(9.5)
    run = par.add_run(testo)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x57, 0x53, 0x14)
    return par


def _elenco(doc, voci, stile="List Bullet", size=10, colore=NERO):
    for voce in voci:
        par = doc.add_paragraph(style=stile)
        run = par.add_run(str(voce))
        run.font.size = Pt(size)
        run.font.color.rgb = colore
        par.paragraph_format.space_after = Pt(2)


def _tabella_campo(doc, campo: dict):
    metriche = campo.get("metriche") or {}
    if not metriche:
        _avviso(doc, "Nessun dato di campo per questo URL.",
                "CrUX non ha traffico sufficiente su questa pagina: restano validi solo i "
                "fatti diagnostici del laboratorio, senza metrica reale né storico.")
        return

    storico = (campo.get("storico") or {}).get("serie", {})
    periodi = (campo.get("storico") or {}).get("periodi", [])

    tabella = doc.add_table(rows=1, cols=5)
    tabella.style = "Table Grid"
    for cella, testo in zip(tabella.rows[0].cells,
                            ["Metrica", "p75 reale", "Giudizio", "Delta 40 sett.", "Andamento"]):
        run = cella.paragraphs[0].add_run(testo)
        run.bold = True
        run.font.size = Pt(8.5)

    for metrica, valore in metriche.items():
        if metrica not in ETICHETTE:
            continue
        g = giudizio(metrica, valore)
        serie = storico.get(metrica, [])
        puliti = [x for x in serie if x is not None]
        delta = ""
        if len(puliti) >= 2:
            d = puliti[-1] - puliti[0]
            delta = f"{d:+.0f} ms" if SOGLIE[metrica].unita else f"{d:+.2f}"

        celle = tabella.add_row().cells
        for indice, testo in enumerate([ETICHETTE[metrica], formatta(metrica, valore),
                                        PAROLA[g], delta, _sparkline(serie)]):
            run = celle[indice].paragraphs[0].add_run(testo)
            run.font.size = Pt(9.5)
            if indice == 0:
                run.bold = True
            if indice == 2:
                run.font.color.rgb = COLORI[g]
                run.bold = True
        celle[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        celle[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if periodi:
        _p(doc, f"Storico {periodi[0]} → {periodi[-1]} · {len(periodi)} settimane · "
               f"delta = ultimo periodo meno primo.", size=8.5, colore=GRIGIO, corsivo=True)


def _fasi_lcp(doc, fatti: dict):
    fasi = fatti.get("lcp_fasi") or {}
    if not fasi:
        return
    totale = sum(fasi.values()) or 1
    dominante = max(fasi, key=fasi.get)

    _p(doc, "Dove si perde il tempo dell'LCP", size=10, grassetto=True, spazio_dopo=6)
    for chiave, valore in fasi.items():
        quota = valore / totale
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(1)
        etichetta = par.add_run(f"{FASI_IT.get(chiave, chiave):<30}")
        etichetta.font.size = Pt(9.5)
        etichetta.font.color.rgb = NERO if chiave == dominante else GRIGIO
        etichetta.bold = chiave == dominante
        barra = par.add_run("█" * max(1, round(quota * 30)) + f"  {quota * 100:.0f}%")
        barra.font.size = Pt(9.5)
        barra.font.color.rgb = COLORI["scarso"] if chiave == dominante else GRIGIO

    snippet = fatti.get("lcp_elemento_snippet") or ""
    if snippet:
        _p(doc, f"Elemento LCP: {snippet[:200]}", size=8.5, colore=GRIGIO, spazio_dopo=2)
    _p(doc, "Proporzioni sul trace osservato: non sommano alla metrica LCP riportata da "
           "Lighthouse, che è simulata con throttling.",
       size=8.5, colore=GRIGIO, corsivo=True, spazio_dopo=10)


FONTE = {
    "lighthouse": "testo di Lighthouse",
    "classificazione": "classificazione su dati di laboratorio",
    "campo": "classificazione su dati di campo",
}


def _risorse(doc, righe: list):
    """I file che causano il problema: risorsa, impatto, prima o terza parte."""
    if not righe:
        return
    tabella = doc.add_table(rows=1, cols=3)
    tabella.style = "Table Grid"
    for cella, testo in zip(tabella.rows[0].cells, ["", "Risorsa", "Impatto"]):
        run = cella.paragraphs[0].add_run(testo)
        run.bold = True
        run.font.size = Pt(8)
    for url, misura, terza in righe:
        celle = tabella.add_row().cells
        for indice, testo in enumerate(["3P" if terza else "1P", url, misura]):
            run = celle[indice].paragraphs[0].add_run(testo)
            run.font.size = Pt(7.5 if indice == 1 else 8)
            if indice == 0:
                run.font.color.rgb = GRIGIO
                run.bold = True
        celle[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _elementi(doc, righe: list):
    """Elementi del DOM: separati dalle risorse perche' un nodo non e' ne' prima
    ne' terza parte, e marcarlo sarebbe un'informazione inventata."""
    if not righe:
        return
    tabella = doc.add_table(rows=1, cols=2)
    tabella.style = "Table Grid"
    for cella, testo in zip(tabella.rows[0].cells, ["Elemento nel DOM", "Impatto"]):
        run = cella.paragraphs[0].add_run(testo)
        run.bold = True
        run.font.size = Pt(8)
    for riferimento, misura, percorso, _snippet in righe:
        celle = tabella.add_row().cells
        run = celle[0].paragraphs[0].add_run(riferimento)
        run.font.size = Pt(7.5)
        if percorso:
            sotto = celle[0].add_paragraph()
            r2 = sotto.add_run(percorso)
            r2.font.size = Pt(6.5)
            r2.font.color.rgb = GRIGIO
        run = celle[1].paragraphs[0].add_run(misura)
        run.font.size = Pt(8)
        celle[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _problemi(doc, problemi: list):
    _p(doc, "Interventi", size=10, grassetto=True, spazio_dopo=6)
    if not problemi:
        _p(doc, "Nessun problema rilevato oltre soglia.", size=10)
        return
    for problema in problemi:
        titolo = problema["titolo"]
        if not problema.get("azionabile", True):
            titolo += "  [non azionabile direttamente]"
        _p(doc, titolo, size=11, grassetto=True,
           colore=GRAVITA.get(problema.get("gravita", "bassa"), GRIGIO), spazio_dopo=1)
        _p(doc, f"Priorità {problema.get('gravita')} · interviene: "
               f"{problema.get('responsabile')} · "
               f"{FONTE.get(problema.get('fonte'), problema.get('fonte'))}",
           size=8.5, colore=GRIGIO, spazio_dopo=4)
        _elenco(doc, problema.get("evidenza", []), size=9, colore=GRIGIO)
        _elenco(doc, problema.get("azioni", []), size=10)
        _risorse(doc, problema.get("risorse") or [])
        _elementi(doc, problema.get("elementi") or [])
        if problema.get("nota"):
            _p(doc, problema["nota"], size=8.5, colore=GRIGIO, corsivo=True, spazio_dopo=2)
        if problema.get("documentazione"):
            _p(doc, f"Documentazione: {problema['documentazione']}",
               size=8, colore=GRIGIO, spazio_dopo=10)
        else:
            _p(doc, "", size=4, spazio_dopo=6)


def _trasversale(doc, pagine: list):
    fasi = {}
    for pagina in pagine:
        f = (pagina.get("fatti") or {}).get("lcp_fasi") or {}
        if f:
            fasi.setdefault(max(f, key=f.get), []).append(pagina["template"])
    if not fasi:
        return
    if len(fasi) == 1:
        fase, template = next(iter(fasi.items()))
        _avviso(doc, "Diagnosi trasversale.",
                f"Tutti i template misurati ({', '.join(template)}) perdono il tempo LCP "
                f"nella stessa fase: {FASI_IT.get(fase, fase).lower()}. Il problema è del "
                f"sito, non delle singole pagine: un intervento solo li sistema tutti.")
    else:
        dettaglio = "; ".join(f"{FASI_IT.get(f, f).lower()}: {', '.join(t)}"
                              for f, t in fasi.items())
        _avviso(doc, "Diagnosi trasversale.",
                f"I template perdono tempo in fasi diverse ({dettaglio}). Sono interventi "
                f"separati: la priorità va data ai template con più traffico.")


def docx_report(esecuzione: dict, percorso):
    doc = Document()
    normale = doc.styles["Normal"]
    normale.font.name = "Calibri"
    normale.font.size = Pt(10)

    doc.add_heading(f"Analisi velocità — {esecuzione.get('cliente')}", level=0)
    _p(doc, f"{esecuzione.get('sito')} · rilevazione del {esecuzione.get('data')} · "
           f"{'mobile' if esecuzione.get('form_factor') == 'PHONE' else 'desktop'} · "
           f"un URL rappresentativo per template",
       size=9.5, colore=GRIGIO, spazio_dopo=12)

    pagine = esecuzione.get("pagine", [])
    _trasversale(doc, pagine)
    _avviso(doc, "Come leggere i numeri.",
            "Le metriche «p75 reale» vengono da CrUX: sono l'esperienza degli utenti veri, "
            "su una finestra mobile di 28 giorni. Un intervento messo online oggi entra in "
            "questi numeri gradualmente e si legge pulito solo dopo quattro settimane. I "
            "fatti diagnostici (elemento LCP, fasi, peso) vengono invece da una misurazione "
            "di laboratorio, che serve a capire perché, non quanto.")

    for indice, pagina in enumerate(pagine):
        if indice:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        doc.add_heading(pagina["template"], level=1)
        _p(doc, pagina["url"], size=8.5, colore=GRIGIO, spazio_dopo=2)
        if pagina.get("consenso"):
            _p(doc, f"{pagina.get('misurazioni', 1)} misurazioni di laboratorio · "
                   f"{pagina['consenso']}", size=8, colore=GRIGIO, corsivo=True, spazio_dopo=8)

        if pagina.get("errore"):
            _avviso(doc, "Misurazione fallita.", str(pagina["errore"]))
            continue

        _tabella_campo(doc, pagina.get("campo") or {})
        doc.add_paragraph()
        _fasi_lcp(doc, pagina.get("fatti") or {})

        terze = pagina.get("terze_parti") or {}
        quota = terze.get("byte_terzi", 0) / (terze.get("byte_totali") or 1)
        _p(doc, "Peso della pagina", size=10, grassetto=True, spazio_dopo=4)
        _p(doc, f"{terze.get('byte_totali', 0) / 1024:.0f} KB su "
               f"{terze.get('richieste_totali', 0)} richieste, di cui "
               f"{terze.get('byte_terzi', 0) / 1024:.0f} KB di terze parti "
               f"({quota * 100:.0f}%).", size=10, spazio_dopo=2)
        per_tipo = pagina.get("peso_per_tipo") or {}
        if per_tipo:
            _p(doc, " · ".join(f"{etichetta_tipo(t)} {b / 1024:.0f} KB"
                               for t, b in list(per_tipo.items())[:6] if b),
               size=8.5, colore=GRIGIO, spazio_dopo=10)

        _problemi(doc, pagina.get("problemi") or [])

    doc.add_paragraph()
    vetrina = ", ".join(
        f"{p['template']} {(p.get('fatti') or {}).get('performance_score')}"
        for p in pagine if (p.get("fatti") or {}).get("performance_score") is not None)
    _p(doc, f"Punteggi PageSpeed Insights al momento della rilevazione: {vetrina or 'n/d'}. "
           f"Sono riportati solo come riferimento: variano fra due misurazioni identiche e "
           f"non vengono usati per nessuna valutazione in questo documento. "
           f"Generato il {date.today():%d/%m/%Y} con speed-audit.",
       size=8.5, colore=GRIGIO, corsivo=True)

    doc.save(str(percorso))
    return percorso
